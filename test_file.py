from asyncio.windows_events import NULL
from posixpath import split
import re
import string
import psycopg2
import docx
import sqlite3
import os

#извлечение симптомов из предобработанного текста
def take_symptoms():
    conn = sqlite3.connect('SyndromesSymptoms.db')
    cursor = conn.cursor()
    conn1 = sqlite3.connect('mkb10.db')
    cursor1 = conn1.cursor()
    regexp=other_patterns[2]['name']
    lst_symp=[]
    n=len(lst)-1
    print (n)
    i=0 
    text_symp=""
    id_syndr=""
    while i<n:
        n1=str(lst[i])
        n1=re.sub(r"(\()", "\(", n1)
        n1=re.sub(r"(\))", "\)", n1)
        n2=str(lst[i+1])
        n2=re.sub(r"(\()", "\(", n2)
        n2=re.sub(r"(\))", "\)", n2)
        
        #выбиреам текст м/у двумя названиями синдромов
        
        regex1=other_patterns[4]['name']
        regex1=regex1.replace('n1', n1)
        regex1=regex1.replace('n2', n2)
       #regex2 = r"("+n1+")((.*?)(\s*))*("+n2+")" #7 pattern
        print(regex1)
        matches = re.finditer(regex1, pre_text, re.MULTILINE)
        print("ok")
        for matchNum, match in enumerate(matches, start=1):
        #f matches:
            print("ok")
            symp = "{match}".format(match = match.group())
            symp = re.sub(lst[i],"",symp)
            print("ok")
            symp = re.sub(lst[i+1],"",symp)
            print("ok")
            matches = re.finditer(regexp, symp, re.MULTILINE)
            print("ok")
            for matchNum, match in enumerate(matches, start=1):
                text_symp ="{match}".format(match = match.group())
                print (text_symp)
                #cursor1.execute("select pattern from all_pattern where id=8;")
                #conn1.commit() 
                regex2=other_patterns[5]['name']
                #regex22 = cursor1.fetchone()[0] 
                text_symp = re.sub(regex2,"", text_symp) #8 pattern
                #text_symp = re.sub(r"(ОСНОВНЫЕ СИМПТОМЫ|СИМПТОМЫ)((.*?)(\s*))+(#)","", text_symp) #8 pattern
                print(text_symp)
                if text_symp.find('||'):
                    for s in re.split(r"(\|\|)", text_symp):
                        s.lstrip()
                        if s!=" ":
                            if s!='||':
                                s=s.lstrip()
                                cursor.execute("select id from symptoms_from_teaching_aid where symptom = ?;", (s,))
                                conn.commit()
                                is_exist = cursor.fetchall()
                                if len(is_exist) == 0:
                                    
                                    cursor.execute("insert into symptoms_from_teaching_aid (symptom, reason) values ( ?, ?);", (s,""))
                                    conn.commit()
                                    cursor.execute("select id from symptoms_from_teaching_aid where symptom = (?);", (s,))
                                    conn.commit() 
                                    last_id_symp = cursor.fetchone()[0]
                                    cursor.execute("select id from syndromes_from_teaching_aid where name = ?;", (str(lst[i]),))
                                    conn.commit() 
                                    id_syndr1=cursor.fetchone()[0]
                                    print("-----------------------------------------------")
                                    print(str(lst[i]))
                                    print(id_syndr1)
                                    #id_syndr=0
                                    
                                    #f (len(len1)!=0):
                                        #print(id_syndr)
                                    #else:
                                        #print(str(lst[i]))
                                    cursor.execute("insert into symptoms_of_syndromes_from_teaching_aid (id_syndrom, id_symptom) values (?, ?);", (id_syndr1,last_id_symp))
                                    conn.commit()
                                #else:
                                    #id_sympt=is_exist[0]
                                    #cursor.execute("insert into symptoms_of_syndromes_from_teaching_aid (id_syndrom, id_symptom) values (?, ?)", (id_syndr,id_sympt))
                                    #conn.commit()


            lst_symp.append(text_symp)
        i+=1
    print (n)
    #print (lst[n])
    
    nn=str(lst[n])
    nn=re.sub(r"(\()", "\(", nn)
    nn=re.sub(r"(\))", "\)", nn)
    #cursor1.execute("select pattern from all_pattern where id=9;")
    #conn1.commit() 
    #regex3 = cursor1.fetchone()[0]
    regex3=other_patterns[6]['name']
    #regex3 = r"("+nn+")((.*?)(\s*))*($)" #9 pattern
    #matches = re.search(regex3, pre_text)
    #if matches:
    matches = re.finditer(regex3, pre_text, re.MULTILINE)
    for matchNum, match in enumerate(matches, start=1):
            symp = "{match}".format(match = match.group())
            symp = re.sub(str(lst[n]),"",symp)
            matches = re.finditer(regexp, symp, re.MULTILINE)
            for matchNum, match in enumerate(matches, start=1):
                text_symp ="{match}".format(match = match.group())
                print (text_symp)
                #cursor1.execute("select pattern from all_pattern where id=8;")
                #conn1.commit() 
                #regex22 = cursor1.fetchone()[0] 
                text_symp = re.sub(regex2,"", text_symp) #8 pattern
                #text_symp = re.sub(r"(ОСНОВНЫЕ СИМПТОМЫ|СИМПТОМЫ)((.*?)(\s*))+(#)","", text_symp) #8 pattern
                print(text_symp)
                if text_symp.find('||'):
                    for s in re.split(r"(\|\|)", text_symp):
                        
                        if s!=" ":
                            if s!='||':
                                s = s.lstrip()
                                cursor.execute("insert into symptoms_from_teaching_aid (symptom, reason) values ( ?, ?);", (s,""))
                                conn.commit()
                                cursor.execute("select id from symptoms_from_teaching_aid where symptom = (?);", (s,))
                                conn.commit() 
                                last_id_symp = cursor.fetchone()[0]
                                cursor.execute("select id from syndromes_from_teaching_aid where name = ?;", (str(lst[n]),))
                                conn.commit() 
                                id_syndr = cursor.fetchone()[0]
                                cursor.execute("insert into symptoms_of_syndromes_from_teaching_aid (id_syndrom, id_symptom) values (?, ?);", (id_syndr,last_id_symp))
                                conn.commit()
            
            lst_symp.append(text_symp)
    cursor1.close()
    conn1.close()
    cursor.close()
    conn.close()

if __name__ == "__main__":
    file_path=("D://учёба//4//4000_simptomov_i_sindromov.docx")
    #file_path= ("C://Users//lizav//OneDrive//Рабочий стол//ресурсы//Новая папка//4000_simptomov_i_sindromov.docx")
    doc = docx.Document(file_path)
    med_section = os.path.splitext(os.path.basename(file_path))[0]
    text1 = []
    for paragraph in doc.paragraphs:
        text1.append(paragraph.text)
    lst_syndroms=[]

    #получение всех заголовков
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                if paragraph.text not in lst_syndroms:
                    lst_syndroms.append(paragraph.text)
                #print(paragraph.text)
    #print(lst_syndroms)
    f = open('xyz.txt','w')  # открытие в режиме записи
    for l in lst_syndroms:
        f.write(l)  

    f.close() 