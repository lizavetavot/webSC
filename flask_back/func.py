import re
import docx
import sqlite3
import os
import sys
from pathlib import Path

ontoPath = str(Path(__file__).parent.joinpath("..", "scivi.onto-master").resolve()) 
sys.path.insert(1, ontoPath)

from onto import Onto

data=Onto.load_from_file("SC_ont1.ont")
ddata=Onto.nodes(data)
#for d in ddata:
   #print(d['id'])
pattern_types=Onto.get_nodes_linked_to(data,Onto.get_nodes_by_name(data,'Паттерн')[0],'is_a')
pre_text_pattern=Onto.get_nodes_linked_to(data,pattern_types[0],'is_a')
other_patterns=Onto.get_nodes_linked_to(data,pattern_types[1],'is_a')
print(pre_text_pattern)
print()
print(other_patterns)


def delete_bd():
    conn = sqlite3.connect('db/SyndromesSymptoms.db')
    cursor = conn.cursor()
    cursor.execute("DELETE from syndromes_from_teaching_aid")
    conn.commit() 
    cursor.execute("DELETE from symptoms_of_syndromes_from_teaching_aid")
    conn.commit() 
    cursor.execute("DELETE from symptoms_from_teaching_aid")
    conn.commit() 
    cursor.close()
    conn.close()

def getcodemkb():
    conn = sqlite3.connect('mkb10.db')
    cursor = conn.cursor()
    cursor.execute("select mkb_code, mkb_name from mkb_etalon where mkb_code!='';")
    conn.commit() 
    rows = cursor.fetchall()
    cursor.close()
    conn.close()
    return rows

def setcodemkb(code, name):
    conn = sqlite3.connect('db/SyndromesSymptoms.db')
    cursor = conn.cursor()
    cursor.execute("update syndromes_from_teaching_aid set expert_code_icd = (?) where name =(?);", (code,name))
    conn.commit()
    cursor.close()
    conn.close()

#заполнение treeview
def  first():
    conn = sqlite3.connect('db/SyndromesSymptoms.db')
    cursor = conn.cursor()
    treeData={}
    cursor.execute("select id, name, expert_code_icd from syndromes_from_teaching_aid;")
    conn.commit() 
    rows = cursor.fetchall()
    for row in rows:
        treeData[row[0]]={'label':row[1], 'nodes':{}, 'index':row[0], 'key0': row[2]}
        cursor.execute("SELECT symptoms_from_teaching_aid.symptom FROM symptoms_from_teaching_aid join symptoms_of_syndromes_from_teaching_aid ON symptoms_from_teaching_aid.id = symptoms_of_syndromes_from_teaching_aid.id_symptom join syndromes_from_teaching_aid ON symptoms_of_syndromes_from_teaching_aid.id_syndrom = syndromes_from_teaching_aid.id where syndromes_from_teaching_aid.name = ?;", (row[1],))
        conn.commit()
        rows2 = cursor.fetchall()
        
        for row2 in rows2:
            treeData[row[0]]['nodes'][row2[0]]={'label':row2[0], 'nodes':{}, 'index':0}

    cursor.close()
    conn.close()
    return treeData

def createnode(treeData, expert_code, index, name):
    conn = sqlite3.connect('db/SyndromesSymptoms.db')
    cursor = conn.cursor()
    
    if not(expert_code in treeData):
        treeData[expert_code]={'label':expert_code, 'nodes':{}, 'index':index}
    treeData[expert_code]["nodes"][name]={'label':name, 'nodes':{}, 'index':index}
    cursor.execute("SELECT symptoms_from_teaching_aid.symptom FROM symptoms_from_teaching_aid join symptoms_of_syndromes_from_teaching_aid ON symptoms_from_teaching_aid.id = symptoms_of_syndromes_from_teaching_aid.id_symptom join syndromes_from_teaching_aid ON symptoms_of_syndromes_from_teaching_aid.id_syndrom = syndromes_from_teaching_aid.id where syndromes_from_teaching_aid.name = ?;", (name,))
    conn.commit()
    rows2 = cursor.fetchall()
    for row2 in rows2:
        treeData[expert_code]["nodes"][name]['nodes'][row2[0]]={'label':row2[0], 'nodes':{}, 'index':0}
    
    cursor.close()
    conn.close()

def second():
    conn = sqlite3.connect('db/SyndromesSymptoms.db')
    cursor = conn.cursor()
    treeData={}
    cursor.execute("select * from syndromes_from_teaching_aid;")
    conn.commit() 
    rows = cursor.fetchall()
    for row in rows:
        #print(row[2])
        #id	med_section	name	reason	code_icd	expert_code_icd
        if (row[5]==''):
            #print(row[5])
            createnode(treeData, "Код не установлен", 0 ,row[2] )
        else:
            conn1 = sqlite3.connect('mkb10.db')
            cursor1 = conn1.cursor()
            cursor1.execute("SELECT mkb_name FROM mkb_etalon where mkb_code = ?;", (row[5],))
            conn1.commit()
            roww = cursor1.fetchall()
            code=row[5]+' '+str(roww)[3:-4]
            createnode(treeData, code, 0 ,row[2] )
            cursor1.close()
            conn1.close()
    cursor.close()
    conn.close()
    return treeData

    

#извлечение симптомов из предобработанного текста
def take_symptoms():
    conn = sqlite3.connect('db/SyndromesSymptoms.db')
    cursor = conn.cursor()
    conn1 = sqlite3.connect('mkb10.db')
    cursor1 = conn1.cursor()
    #cursor1.execute("select pattern from all_pattern where id=5;")
    #conn1.commit() 
    #regexp = cursor1.fetchone()[0] 
    regexp=other_patterns[2]['name']
    lst_symp=[]
    n=len(lst)-1
    print (n)
    i=0 
    text_symp=""
    id_syndr=""
    while i<n:
        n1=str(lst[i])
        n1=re.sub(r"(\()", "\(", n1)
        n1=re.sub(r"(\))", "\)", n1)
        n2=str(lst[i+1])
        n2=re.sub(r"(\()", "\(", n2)
        n2=re.sub(r"(\))", "\)", n2)
        
        #выбиреам текст м/у двумя названиями синдромов
        #cursor1.execute("select pattern from all_pattern where id=7;")
        #conn1.commit() 
        #regex2 = cursor1.fetchone()[0] 
        regex1=other_patterns[4]['name']
        regex1=regex1.replace('n1', n1)
        regex1=regex1.replace('n2', n2)
       #regex2 = r"("+n1+")((.*?)(\s*))*("+n2+")" #7 pattern
        print(regex1)
        matches = re.finditer(regex1, pre_text, re.MULTILINE)
        print("ok")
        for matchNum, match in enumerate(matches, start=1):
        #f matches:
            print("ok")
            symp = "{match}".format(match = match.group())
            symp = re.sub(lst[i],"",symp)
            print("ok")
            symp = re.sub(lst[i+1],"",symp)
            print("ok")
            matches = re.finditer(regexp, symp, re.MULTILINE)
            print("ok")
            for matchNum, match in enumerate(matches, start=1):
                text_symp ="{match}".format(match = match.group())
                print (text_symp)
                #cursor1.execute("select pattern from all_pattern where id=8;")
                #conn1.commit() 
                regex2=other_patterns[5]['name']
                #regex22 = cursor1.fetchone()[0] 
                text_symp = re.sub(regex2,"", text_symp) #8 pattern
                #text_symp = re.sub(r"(ОСНОВНЫЕ СИМПТОМЫ|СИМПТОМЫ)((.*?)(\s*))+(#)","", text_symp) #8 pattern
                print(text_symp)
                if text_symp.find('||'):
                    for s in re.split(r"(\|\|)", text_symp):
                        s.lstrip()
                        if s!=" ":
                            if s!='||':
                                s=s.lstrip()
                                cursor.execute("select id from symptoms_from_teaching_aid where symptom = ?;", (s,))
                                conn.commit()
                                is_exist = cursor.fetchall()
                                if len(is_exist) == 0:
                                    
                                    cursor.execute("insert into symptoms_from_teaching_aid (symptom, reason) values ( ?, ?);", (s,""))
                                    conn.commit()
                                    cursor.execute("select id from symptoms_from_teaching_aid where symptom = (?);", (s,))
                                    conn.commit() 
                                    last_id_symp = cursor.fetchone()[0]
                                    cursor.execute("select id from syndromes_from_teaching_aid where name = ?;", (str(lst[i]),))
                                    conn.commit() 
                                    id_syndr1=cursor.fetchone()[0]
                                    print("-----------------------------------------------")
                                    print(str(lst[i]))
                                    print(id_syndr1)
                                    #id_syndr=0
                                    
                                    #f (len(len1)!=0):
                                        #print(id_syndr)
                                    #else:
                                        #print(str(lst[i]))
                                    cursor.execute("insert into symptoms_of_syndromes_from_teaching_aid (id_syndrom, id_symptom) values (?, ?);", (id_syndr1,last_id_symp))
                                    conn.commit()
                                #else:
                                    #id_sympt=is_exist[0]
                                    #cursor.execute("insert into symptoms_of_syndromes_from_teaching_aid (id_syndrom, id_symptom) values (?, ?)", (id_syndr,id_sympt))
                                    #conn.commit()


            lst_symp.append(text_symp)
        i+=1
    print (n)
    #print (lst[n])
    
    nn=str(lst[n])
    nn=re.sub(r"(\()", "\(", nn)
    nn=re.sub(r"(\))", "\)", nn)
    #cursor1.execute("select pattern from all_pattern where id=9;")
    #conn1.commit() 
    #regex3 = cursor1.fetchone()[0]
    regex3=other_patterns[6]['name']
    #regex3 = r"("+nn+")((.*?)(\s*))*($)" #9 pattern
    #matches = re.search(regex3, pre_text)
    #if matches:
    matches = re.finditer(regex3, pre_text, re.MULTILINE)
    for matchNum, match in enumerate(matches, start=1):
            symp = "{match}".format(match = match.group())
            symp = re.sub(str(lst[n]),"",symp)
            matches = re.finditer(regexp, symp, re.MULTILINE)
            for matchNum, match in enumerate(matches, start=1):
                text_symp ="{match}".format(match = match.group())
                print (text_symp)
                #cursor1.execute("select pattern from all_pattern where id=8;")
                #conn1.commit() 
                #regex22 = cursor1.fetchone()[0] 
                text_symp = re.sub(regex2,"", text_symp) #8 pattern
                #text_symp = re.sub(r"(ОСНОВНЫЕ СИМПТОМЫ|СИМПТОМЫ)((.*?)(\s*))+(#)","", text_symp) #8 pattern
                print(text_symp)
                if text_symp.find('||'):
                    for s in re.split(r"(\|\|)", text_symp):
                        
                        if s!=" ":
                            if s!='||':
                                s = s.lstrip()
                                cursor.execute("insert into symptoms_from_teaching_aid (symptom, reason) values ( ?, ?);", (s,""))
                                conn.commit()
                                cursor.execute("select id from symptoms_from_teaching_aid where symptom = (?);", (s,))
                                conn.commit() 
                                last_id_symp = cursor.fetchone()[0]
                                cursor.execute("select id from syndromes_from_teaching_aid where name = ?;", (str(lst[n]),))
                                conn.commit() 
                                id_syndr = cursor.fetchone()[0]
                                cursor.execute("insert into symptoms_of_syndromes_from_teaching_aid (id_syndrom, id_symptom) values (?, ?);", (id_syndr,last_id_symp))
                                conn.commit()
            
            lst_symp.append(text_symp)
    cursor1.close()
    conn1.close()
    cursor.close()
    conn.close()

#предобработка текста
def make_preprocessing(text):
    while '\n' in text:
        text= text.replace('\n', " ")
    
    text = " ".join(text.split())
    
    #text= re.sub("\s+", "\s",text)
    while '' in text:
            text= text.replace('', '-')
    
    while ")" in text:
            text= text.replace(')', ')')
  
    while ")•" in text:
            text= text.replace(')•', ')')
    
    while "" in text:
            text= text.replace('', '##')
    
    while "•" in text:
            text= text.replace('•', '##')
    conn1 = sqlite3.connect('mkb10.db')
    cursor1 = conn1.cursor()
    #cursor1.execute("select pattern from all_pattern where id=6;")
    #conn1.commit() 
    #global regex
    #regex = cursor1.fetchone()[0] 
    regex4=other_patterns[3]['name']
    #cursor1.execute("select pattern from all_pattern where id=4;")
    #conn1.commit() 

    #global regex1 #переименовать
    #regex1 = cursor1.fetchone()[0]    
    regex5=pre_text_pattern[0]['name'] 
    # замена : после слова СИМПТОМЫ в перечислении симптомов на #
    matches = re.finditer(regex5, text, re.MULTILINE)
    for matchNum, match in enumerate(matches, start=1):
            s="{match}".format(match = match.group())
            new_s= re.sub(":", '#', s)
            text=text.replace(s, new_s)
    #замена мааркеров нумерованного списка, знако препинания при перечислении симптомов
    matches = re.finditer(regex4, text, re.MULTILINE)
    for matchNum, match in enumerate(matches, start=1):
            s="{match}".format(match = match.group())
            new_s= re.sub(r"(\s)([0-9]+\))", ' ||', s)
            new_s = re.sub(",", '|', new_s)
            new_s = re.sub(";", '|', new_s)
            text=text.replace(s, new_s)
    return text
    cursor1.close()
    conn1.close()

def _add(file_path):
    conn = sqlite3.connect('db/SyndromesSymptoms.db')
    cursor = conn.cursor()
    file_path= re.sub ("/", "//", file_path)
    doc = docx.Document(file_path)
    med_section = os.path.splitext(os.path.basename(file_path))[0]
    text1 = []
    for paragraph in doc.paragraphs:
        text1.append(paragraph.text)
        
    global pre_text
    pre_text = '\n'.join(text1)
    pre_text = make_preprocessing(pre_text)
    global lst1
    lst1 =[]
    global lst
    lst = []
    #получение всех заголовков
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                lst1.append(paragraph.text)
                #print(paragraph.text)
        
    for item1 in lst1: 
        if item1 not in lst: 
            lst.append(item1) 
    for item in lst:
        
        #заполнение таблицы БД синдромами
        cursor.execute("select name from syndromes_from_teaching_aid where name = (?);", (item,))
        conn.commit()
        is_exist = cursor.fetchall()
        if len(is_exist) == 0:
            cursor.execute("insert into syndromes_from_teaching_aid (med_section, name, reason, code_icd, expert_code_icd) values (?, ?, ?, ?, ?);", (med_section, item,"","", ""))
            conn.commit()
    take_symptoms()
    cursor.close()
    conn.close()


#D:\\учёба\\3курс\\практика\\SymptCheckerSQLite\\
conn1 = sqlite3.connect('mkb10.db')
#conn1 = sqlite3.connect('D:\\учёба\\3курс\\практика\\SymptCheckerSQLite\\mkb10.db')
cursor1 = conn1.cursor()

conn = sqlite3.connect('db/SyndromesSymptoms.db')
cursor = conn.cursor()
#conn = psycopg2.connect(dbname='Diseases_and_Symptoms', user='postgres', password = 'polina')
#cursor = conn.cursor()
cursor.execute("""CREATE TABLE IF NOT EXISTS syndromes_from_teaching_aid(
id INTEGER PRIMARY KEY AUTOINCREMENT,
med_section TEXT,
name TEXT,
reason TEXT,
code_icd TEXT,
expert_code_icd TEXT);
""")
conn.commit() 

cursor.execute("""CREATE TABLE IF NOT EXISTS symptoms_from_teaching_aid(
id INTEGER PRIMARY KEY AUTOINCREMENT,
symptom TEXT,
reason TEXT);
""")
conn.commit() 

cursor.execute("""CREATE TABLE IF NOT EXISTS symptoms_of_syndromes_from_teaching_aid(
id_syndrom INT,
id_symptom INT);
""")
conn.commit()   

lst=[]
list_dis=[]
lst_code=[]
lst_name=[]

cursor.close()
conn.close()
cursor1.close()
conn1.close()

